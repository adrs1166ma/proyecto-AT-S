import io
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from fpdf import FPDF

MESES = {
    1: "enero", 2: "febrero", 3: "marzo", 4: "abril",
    5: "mayo", 6: "junio", 7: "julio", 8: "agosto",
    9: "setiembre", 10: "octubre", 11: "noviembre", 12: "diciembre"
}

AGENTES = {
    "PQS":              "PQS POLVOS QUIMICOS SECO",
    "CO2":              "GAS CARBONICO (CO2)",
    "GAS_PRESURIZADA":  "GAS PRESURIZADA",
    "ACETATO_POTASIO":  "ACETATO DE POTASIO",
}


def _fecha_larga(s: str) -> str:
    d = datetime.strptime(s, "%Y-%m-%d")
    return f"Lima, {d.day:02d} de {MESES[d.month]} del {d.year}"

def _fecha_corta(s: str) -> str:
    d = datetime.strptime(s, "%Y-%m-%d")
    return f"{d.day:02d}-{d.month:02d}-{d.year}"


# ================================================================
# DOCX
# ================================================================

def generar_docx(req) -> bytes:
    doc = Document()

    for sec in doc.sections:
        sec.top_margin    = Cm(1.5)
        sec.bottom_margin = Cm(1.5)
        sec.left_margin   = Cm(2)
        sec.right_margin  = Cm(2)

    def par(text="", bold=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
        p = doc.add_paragraph()
        p.alignment = align
        if text:
            r = p.add_run(text)
            r.bold   = bold
            r.italic = italic
            r.font.size = Pt(size)
        return p

    def campo(label, valor):
        p = doc.add_paragraph()
        r1 = p.add_run(f"{label} : ")
        r1.bold = True
        p.add_run(valor)

    # --- Encabezado ---
    par("AT&S INVERSIONES ANTASHELY S.A.C", bold=True, size=14,
        align=WD_ALIGN_PARAGRAPH.CENTER)
    par("Extintores y fumigaciones", italic=True,
        align=WD_ALIGN_PARAGRAPH.CENTER)
    par()

    # --- Datos cliente ---
    label_nombre = "NOMBRE O RAZÓN SOCIAL" if req.cliente.ruc else "NOMBRE"
    campo(label_nombre, req.cliente.nombre.upper())
    if req.cliente.ruc:
        campo("RUC", req.cliente.ruc)
    campo("Dirección", req.cliente.direccion)
    campo("DISTRITO", req.cliente.distrito.upper())
    par()

    # --- Título ---
    par("CERTIFICADO", bold=True, size=20, align=WD_ALIGN_PARAGRAPH.CENTER)
    tipo_txt = "RECARGA DE EXTINTOR" if req.tipo == "RECARGA" else "EXTINTOR NUEVO"
    par(tipo_txt, bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
    par()

    # --- Cuerpo ---
    agente  = AGENTES.get(req.tipo_agente, req.tipo_agente)
    accion  = "RECARGA" if req.tipo == "RECARGA" else "INSTALACIÓN"
    cuerpo  = (
        f"Se certifica haber realizado la {accion} DE EXTINTORES, conforme a lo "
        f"establecido en la Norma Técnica Peruana NTP 350.043-1/2011, Garantizando "
        f"la operatividad al 100% de los siguientes equipos contra incendio con {agente}."
    )
    par(cuerpo)

    if req.prueba_hidrostatica and req.fecha_prueba_hidrostatica:
        par(f"PRUEBA HIDROSTATICA - Fecha: {_fecha_corta(req.fecha_prueba_hidrostatica)}")
    par()

    # --- Tabla extintores ---
    tiene_marca = any(e.marca for e in req.extintores)
    tiene_serie = any(e.serie for e in req.extintores)

    encabezados = ["N.", "CAPAC", "CLASE"]
    if tiene_marca: encabezados.append("MARCA")
    if tiene_serie: encabezados.append("SERIE")
    encabezados += ["FECHA RECARGA", "FECHA VENC."]

    tabla = doc.add_table(rows=1, cols=len(encabezados))
    tabla.style = "Table Grid"
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(encabezados):
        c = tabla.rows[0].cells[i]
        c.text = h
        c.paragraphs[0].runs[0].bold = True

    for idx, ext in enumerate(req.extintores, 1):
        fila = tabla.add_row().cells
        vals = [str(idx), ext.capacidad, ext.clase]
        if tiene_marca: vals.append(ext.marca or "")
        if tiene_serie: vals.append(ext.serie or "")
        vals += [_fecha_corta(ext.fecha_recarga), _fecha_corta(ext.fecha_vencimiento)]
        for i, v in enumerate(vals):
            fila[i].text = v

    par()

    # --- Garantía ---
    garantia = (
        "Esta garantía quedara sin efecto en caso que el equipo haya sido utilizado. "
        "El precinto de seguridad haya sido roto, manipulado y/o trasladado a lugares "
        "inadecuadas (áreas de temperatura sobredimensionadas a su normalidad) por "
        "terceros no autorizados por:"
    )
    p_g = doc.add_paragraph()
    p_g.add_run(garantia).bold = True
    p_g = doc.add_paragraph()
    p_g.add_run("AT&S INVERSIONES ANTASHELY S.A.C").bold = True

    par()
    fecha_p = doc.add_paragraph()
    fecha_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fecha_p.add_run(_fecha_larga(req.fecha_emision))

    par()
    par("AT&S INVERSIONES ANTASHELY S.A.C", bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER)
    par("Gerente Comercial", align=WD_ALIGN_PARAGRAPH.CENTER)
    par("Alfredo Mancilla", align=WD_ALIGN_PARAGRAPH.CENTER)
    par()
    par("Villa el salvador - Lima  |  antashely@hotmail.com  |  960 183 893",
        align=WD_ALIGN_PARAGRAPH.CENTER)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ================================================================
# PDF
# ================================================================

class PDF(FPDF):
    def header(self):
        self.set_font("Helvetica", "B", 13)
        self.cell(0, 8, "AT&S INVERSIONES ANTASHELY S.A.C", align="C", new_x="LMARGIN", new_y="NEXT")
        self.set_font("Helvetica", "I", 9)
        self.cell(0, 5, "Extintores y fumigaciones", align="C", new_x="LMARGIN", new_y="NEXT")
        self.ln(4)

    def footer(self):
        self.set_y(-15)
        self.set_font("Helvetica", "", 8)
        self.cell(0, 5,
            "Villa el salvador - Lima  |  antashely@hotmail.com  |  960 183 893",
            align="C")


def generar_pdf(req) -> bytes:
    pdf = PDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_margins(20, 20, 20)

    def line(text, bold=False, size=10, align="L", ln=True):
        pdf.set_font("Helvetica", "B" if bold else "", size)
        pdf.multi_cell(0, 6, text, align=align, new_x="LMARGIN", new_y="NEXT" if ln else "RIGHT")

    def campo_pdf(label, valor):
        pdf.set_font("Helvetica", "B", 10)
        pdf.cell(55, 6, f"{label} :", new_x="RIGHT", new_y="TOP")
        pdf.set_font("Helvetica", "", 10)
        pdf.multi_cell(0, 6, valor, new_x="LMARGIN", new_y="NEXT")

    # --- Datos cliente ---
    label_nombre = "NOMBRE O RAZON SOCIAL" if req.cliente.ruc else "NOMBRE"
    campo_pdf(label_nombre, req.cliente.nombre.upper())
    if req.cliente.ruc:
        campo_pdf("RUC", req.cliente.ruc)
    campo_pdf("Direccion", req.cliente.direccion)
    campo_pdf("DISTRITO", req.cliente.distrito.upper())
    pdf.ln(5)

    # --- Título ---
    line("CERTIFICADO", bold=True, size=18, align="C")
    tipo_txt = "RECARGA DE EXTINTOR" if req.tipo == "RECARGA" else "EXTINTOR NUEVO"
    line(tipo_txt, bold=True, size=12, align="C")
    pdf.ln(4)

    # --- Cuerpo ---
    agente = AGENTES.get(req.tipo_agente, req.tipo_agente)
    accion = "RECARGA" if req.tipo == "RECARGA" else "INSTALACION"
    cuerpo = (
        f"Se certifica haber realizado la {accion} DE EXTINTORES, conforme a lo "
        f"establecido en la Norma Tecnica Peruana NTP 350.043-1/2011, Garantizando "
        f"la operatividad al 100% de los siguientes equipos contra incendio con {agente}."
    )
    line(cuerpo, size=9)
    if req.prueba_hidrostatica and req.fecha_prueba_hidrostatica:
        line(f"PRUEBA HIDROSTATICA - Fecha: {_fecha_corta(req.fecha_prueba_hidrostatica)}", size=9)
    pdf.ln(4)

    # --- Tabla extintores ---
    tiene_marca = any(e.marca for e in req.extintores)
    tiene_serie = any(e.serie for e in req.extintores)

    encabezados = ["N.", "CAPAC", "CLASE"]
    if tiene_marca: encabezados.append("MARCA")
    if tiene_serie: encabezados.append("SERIE")
    encabezados += ["FECHA RECARGA", "FECHA VENC."]

    ancho_total = 170
    n_cols = len(encabezados)
    col_w = [10, 22, 18]
    if tiene_marca: col_w.append(25)
    if tiene_serie: col_w.append(20)
    restante = ancho_total - sum(col_w)
    col_w += [restante // 2, restante - restante // 2]

    pdf.set_font("Helvetica", "B", 9)
    pdf.set_fill_color(220, 220, 220)
    for i, h in enumerate(encabezados):
        pdf.cell(col_w[i], 7, h, border=1, fill=True, align="C")
    pdf.ln()

    pdf.set_font("Helvetica", "", 9)
    for idx, ext in enumerate(req.extintores, 1):
        vals = [str(idx), ext.capacidad, ext.clase]
        if tiene_marca: vals.append(ext.marca or "")
        if tiene_serie: vals.append(ext.serie or "")
        vals += [_fecha_corta(ext.fecha_recarga), _fecha_corta(ext.fecha_vencimiento)]
        for i, v in enumerate(vals):
            pdf.cell(col_w[i], 7, v, border=1, align="C")
        pdf.ln()

    pdf.ln(5)

    # --- Garantía ---
    garantia = (
        "Esta garantia quedara sin efecto en caso que el equipo haya sido utilizado. "
        "El precinto de seguridad haya sido roto, manipulado y/o trasladado a lugares "
        "inadecuadas por terceros no autorizados por:"
    )
    pdf.set_font("Helvetica", "B", 9)
    pdf.multi_cell(0, 5, garantia, new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "B", 10)
    pdf.cell(0, 6, "AT&S INVERSIONES ANTASHELY S.A.C", new_x="LMARGIN", new_y="NEXT")

    pdf.ln(4)
    pdf.set_font("Helvetica", "", 10)
    pdf.cell(0, 6, _fecha_larga(req.fecha_emision), align="R", new_x="LMARGIN", new_y="NEXT")

    pdf.ln(12)
    pdf.set_font("Helvetica", "B", 10)
    pdf.cell(0, 6, "AT&S INVERSIONES ANTASHELY S.A.C", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 10)
    pdf.cell(0, 6, "Gerente Comercial", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.cell(0, 6, "Alfredo Mancilla", align="C", new_x="LMARGIN", new_y="NEXT")

    return bytes(pdf.output())
